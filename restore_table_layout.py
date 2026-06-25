import os
import copy
import json
import argparse
import logging
from docx import Document

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger("layout_restoration")

def copy_xml_properties(src_element, tgt_element, pr_name):
    """Safely copies XML properties element by deepcopying."""
    src_pr = getattr(src_element, pr_name, None)
    if src_pr is not None:
        tgt_pr = getattr(tgt_element, pr_name, None)
        if tgt_pr is not None:
            try:
                idx = tgt_element.index(tgt_pr)
                tgt_element.remove(tgt_pr)
                tgt_element.insert(idx, copy.deepcopy(src_pr))
            except ValueError:
                tgt_element.insert(0, copy.deepcopy(src_pr))
        else:
            if pr_name == "tblGrid":
                tblPr = getattr(tgt_element, "tblPr", None)
                if tblPr is not None:
                    try:
                        idx = tgt_element.index(tblPr)
                        tgt_element.insert(idx + 1, copy.deepcopy(src_pr))
                    except ValueError:
                        tgt_element.insert(0, copy.deepcopy(src_pr))
                else:
                    tgt_element.insert(0, copy.deepcopy(src_pr))
            else:
                tgt_element.insert(0, copy.deepcopy(src_pr))

def restore_layouts(src_doc_path, tgt_doc_path, output_doc_path, test_mode=False):
    logger.info(f"Opening original document: {src_doc_path}")
    src_doc = Document(src_doc_path)
    
    logger.info(f"Opening localized document: {tgt_doc_path}")
    tgt_doc = Document(tgt_doc_path)
    
    report = {
        "tables_processed": 0,
        "tables_skipped": 0,
        "mismatches": [],
        "errors": []
    }
    
    num_tables = min(len(src_doc.tables), len(tgt_doc.tables))
    if test_mode:
        logger.info("Test layout mode enabled. Processing only the first table.")
        num_tables = min(1, num_tables)
        
    for t_idx in range(num_tables):
        src_table = src_doc.tables[t_idx]
        tgt_table = tgt_doc.tables[t_idx]
        
        # Verify row and column matching
        src_rows_len = len(src_table.rows)
        tgt_rows_len = len(tgt_table.rows)
        
        # We also inspect col count from first row
        src_cols_len = len(src_table.rows[0].cells) if src_rows_len > 0 else 0
        tgt_cols_len = len(tgt_table.rows[0].cells) if tgt_rows_len > 0 else 0
        
        if src_rows_len != tgt_rows_len or src_cols_len != tgt_cols_len:
            msg = f"Table {t_idx} structural mismatch: Original has ({src_rows_len}x{src_cols_len}), Localized has ({tgt_rows_len}x{tgt_cols_len}). Skipping."
            logger.warning(msg)
            report["mismatches"].append({
                "table_index": t_idx,
                "reason": msg
            })
            report["tables_skipped"] += 1
            continue
            
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            
            # Ensure table layout is set to fixed to prevent auto-fit behavior on Thai text
            tblPr = tgt_table._tbl.tblPr
            tblLayout = tblPr.first_child_found_in("w:tblLayout")
            if tblLayout is not None:
                tblPr.remove(tblLayout)
            tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))
            
            # 1. Copy table-level properties (tblPr)
            copy_xml_properties(src_table._tbl, tgt_table._tbl, "tblPr")
            
            # Copy table column widths definition (tblGrid)
            copy_xml_properties(src_table._tbl, tgt_table._tbl, "tblGrid")
            

                
            # Copy column widths definition
            for col_idx in range(min(len(src_table.columns), len(tgt_table.columns))):
                tgt_table.columns[col_idx].width = src_table.columns[col_idx].width
            
            # 2. Copy row-level properties (trPr) and cell-level properties (tcPr)
            for r_idx in range(len(src_table.rows)):
                src_row = src_table.rows[r_idx]
                tgt_row = tgt_table.rows[r_idx]
                
                # Copy row properties
                copy_xml_properties(src_row._tr, tgt_row._tr, "trPr")
                if src_row.height:
                    tgt_row.height = src_row.height
                
                # Copy cell properties
                for c_idx in range(len(src_row.cells)):
                    src_cell = src_row.cells[c_idx]
                    tgt_cell = tgt_row.cells[c_idx]
                    
                    # Store original text to verify no text changes occurred
                    original_txt = tgt_cell.text
                    
                    # Copy cell width property
                    if src_cell.width:
                        tgt_cell.width = src_cell.width
                    
                    copy_xml_properties(src_cell._tc, tgt_cell._tc, "tcPr")
                    
                    # Verify text remains unchanged
                    if tgt_cell.text != original_txt:
                        # If XML copying somehow affected paragraphs (should not, as tcPr is a child element of tc,
                        # not containing paragraphs), log a mismatch error
                        logger.error(f"Text mismatch detected in Table {t_idx}, Row {r_idx}, Cell {c_idx}")
                        report["errors"].append({
                            "table_index": t_idx,
                            "row_index": r_idx,
                            "cell_index": c_idx,
                            "error": "Text contents were modified during layout restore"
                        })
            
            report["tables_processed"] += 1
            logger.info(f"Successfully restored table layout for Table {t_idx}.")
            
        except Exception as e:
            logger.error(f"Failed to restore layout for Table {t_idx}: {e}")
            report["errors"].append({
                "table_index": t_idx,
                "error": str(e)
            })
            report["tables_skipped"] += 1
            
    # Save the output
    logger.info(f"Saving fixed layout document to: {output_doc_path}")
    tgt_doc.save(output_doc_path)
    
    # Save report
    report_path = "layout_restore_report.json"
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2)
    logger.info(f"Restoration report saved to {report_path}")
    return report

def main():
    parser = argparse.ArgumentParser(description="DOCX Table Layout Restoration Step")
    parser.add_argument("--src", default="SMENFP1.docx", help="Original template file")
    parser.add_argument("--tgt", default="SMENFP1_th.docx", help="Localized file to apply layouts to")
    parser.add_argument("--out", default="SMENFP1_th_fixed.docx", help="Output fixed file")
    parser.add_argument("--test-layout", action="store_true", help="Process only the first table to test layouts")
    
    args = parser.parse_args()
    
    restore_layouts(args.src, args.tgt, args.out, test_mode=args.test_layout)

if __name__ == "__main__":
    main()
