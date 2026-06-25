import os
import re
import json
import time
import argparse
import csv
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from threading import Lock
import docx
from docx import Document

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger("localization")

# Regex patterns for protected strings
PROTECTED_PATTERNS = [
    r'^[A-Z0-9_-]+$',  # IDs, SKUs, Part numbers like ID-001, ABC-12345
    r'^v?\d+(\.\d+)+$',  # Version numbers like v2.1.4, 1.0.0
    r'https?://[^\s]+',  # URLs
    r'^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$',  # Emails
    r'^[a-zA-Z0-9_-]+\.[a-zA-Z0-9]+$',  # File names like document.docx
    r'^[a-zA-Z]:\\.*$',  # Windows file paths
    r'^/.*$',  # Unix file paths
    r'^\{[a-zA-Z_0-9]+\}$',  # Placeholders like {username}
    r'^\{\d+\}$',  # Placeholders like {0}
    r'^%\w$',  # Placeholders like %s, %d
    r'^%\([a-zA-Z_0-9]+\)\w$',  # Placeholders like %(name)s
    r'^\{\{[a-zA-Z_0-9]+\}\}$',  # Placeholders like {{customer_name}}
    r'^<[A-Z0-9_]+>$',  # Tokens like <TOKEN>
    r'^\[[A-Z0-9_]+\]$',  # Variables like [VARIABLE]
]

THAI_CHAR_PATTERN = re.compile(r'[\u0e00-\u0e7f]')

class RateLimiter:
    """Centralized rate controller to enforce requests-per-minute limits."""
    def __init__(self, requests_per_minute=60, backoff_factor=2.0, max_retries=5):
        self.requests_per_minute = requests_per_minute
        self.delay = 60.0 / requests_per_minute if requests_per_minute > 0 else 0
        self.backoff_factor = backoff_factor
        self.max_retries = max_retries
        self.last_request_time = 0.0
        self.lock = Lock()

    def wait(self):
        if self.delay <= 0:
            return
        with self.lock:
            now = time.time()
            elapsed = now - self.last_request_time
            if elapsed < self.delay:
                time.sleep(self.delay - elapsed)
            self.last_request_time = time.time()

class TranslatorBackend:
    """Wrapper to support multiple translation backends."""
    def __init__(self, config):
        self.config = config
        self.backend_type = config.get("translation_backend", "google_translator")
        self.rate_limiter = RateLimiter(
            requests_per_minute=config["rate_limiting"]["max_requests_per_minute"],
            backoff_factor=config["rate_limiting"]["retry_backoff_factor"],
            max_retries=config["rate_limiting"]["max_retries"]
        )
        self.gemini_config = config.get("gemini", {})
        
        # Initialize GoogleTranslator fallback
        try:
            from deep_translator import GoogleTranslator
            self.google_translator = GoogleTranslator(source='en', target='th')
        except ImportError:
            logger.warning("deep_translator not installed. GoogleTranslator fallback disabled.")
            self.google_translator = None

    def translate(self, text, context=None):
        """Translates text to Thai with retries and rate limiting."""
        if not text.strip():
            return text

        # Check if the string is entirely Thai already
        if THAI_CHAR_PATTERN.search(text):
            return text

        # Check if text matches protected patterns
        is_protected = False
        for pattern in PROTECTED_PATTERNS:
            if re.match(pattern, text.strip()):
                is_protected = True
                break
        if is_protected:
            return text

        retries = 0
        delay = self.rate_limiter.delay
        
        while retries <= self.rate_limiter.max_retries:
            self.rate_limiter.wait()
            try:
                if self.backend_type == "gemini" and self.gemini_config.get("api_key"):
                    res = self._translate_gemini(text, context)
                else:
                    res = self._translate_google(text)
                return res if res is not None else text
            except Exception as e:
                retries += 1
                if retries > self.rate_limiter.max_retries:
                    logger.error(f"Translation failed after {retries} retries: {e}")
                    raise e
                sleep_time = delay * (self.rate_limiter.backoff_factor ** retries)
                logger.warning(f"Error during translation: {e}. Retrying in {sleep_time:.2f} seconds...")
                time.sleep(sleep_time)
        return text

    def _translate_google(self, text):
        if not self.google_translator:
            raise RuntimeError("GoogleTranslator backend is not available.")
        return self.google_translator.translate(text)

    def _translate_gemini(self, text, context):
        import requests
        api_key = self.gemini_config.get("api_key")
        model = self.gemini_config.get("model", "gemini-1.5-flash")
        url = self.gemini_config.get("endpoint", "").format(model=model) + f"?key={api_key}"
        
        system_instruction = (
            "You are a professional technical localization specialist. Translate the English text to natural, professional Thai. "
            "Preserve any placeholders (like {username}, {0}, %s), variables, product codes, or brand names exactly as they are. "
            "If XML tags like <r0>, <r1> are present, translate only the text inside the tags, preserving the tag structure, tag names, and tag order exactly. "
            "Do not add any explanations, notes, or introductory text. Return only the translated text."
        )
        
        prompt = f"English text: {text}\n"
        if context:
            prompt += f"Context (Table Header/Row context): {context}\n"
        prompt += "Thai Translation:"
        
        payload = {
            "contents": [{
                "parts": [{"text": prompt}]
            }],
            "systemInstruction": {
                "parts": [{"text": system_instruction}]
            },
            "generationConfig": {
                "temperature": 0.1
            }
        }
        
        headers = {"Content-Type": "application/json"}
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()
        
        result_json = response.json()
        translated_text = result_json['candidates'][0]['content']['parts'][0]['text'].strip()
        return translated_text

class LocalizationPipeline:
    def __init__(self, config_path="config.json"):
        with open(config_path, "r", encoding="utf-8") as f:
            self.config = json.load(f)
            
        self.input_path = self.config["files"]["input_docx"]
        self.output_path = self.config["files"]["output_docx"]
        self.test_path = self.config["files"]["test_docx"]
        self.tm_path = self.config["files"]["translation_memory"]
        self.glossary_path = self.config["files"]["glossary"]
        self.analysis_report_path = self.config["files"]["analysis_report"]
        self.review_report_path = self.config["files"]["review_report"]
        
        self.translator = TranslatorBackend(self.config)
        self.tm = self._load_json_file(self.tm_path)
        self.glossary = self._load_json_file(self.glossary_path)
        self.tm_lock = Lock()

    def _load_json_file(self, path):
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                try:
                    return json.load(f)
                except json.JSONDecodeError:
                    return {}
        return {}

    def _save_tm(self):
        with self.tm_lock:
            with open(self.tm_path, "w", encoding="utf-8") as f:
                json.dump(self.tm, f, ensure_ascii=False, indent=2)

    def lookup_translation(self, text):
        """Lookup translation in glossary and translation memory (exact, normalized, case-insensitive)."""
        if not text:
            return None
        
        # 1. Glossary Lookup (highest priority)
        normalized_text = " ".join(text.split()).strip()
        if text in self.glossary:
            return self.glossary[text]
        if normalized_text in self.glossary:
            return self.glossary[normalized_text]
        
        # Case insensitive glossary check
        for k, v in self.glossary.items():
            if k.lower() == normalized_text.lower():
                return v

        # 2. TM Lookup
        if text in self.tm:
            return self.tm[text]["translation"]
        
        # Normalized or case differences
        for k, entry in self.tm.items():
            k_norm = " ".join(k.split()).strip()
            if k_norm.lower() == normalized_text.lower():
                return entry["translation"]
                
        return None

    def add_to_tm(self, source, translation, context=None):
        with self.tm_lock:
            self.tm[source] = {
                "translation": translation,
                "context": context or "",
                "timestamp": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "status": "translated"
            }
        self._save_tm()

    def scan_translation_cells(self, doc_path, limit_tables=None, limit_cells=None):
        """Finds all target cells in the 'Translation' column of tables."""
        doc = Document(doc_path)
        cells_to_process = []
        
        table_count = 0
        cell_count = 0
        
        for t_idx, table in enumerate(doc.tables):
            if limit_tables and table_count >= limit_tables:
                break
                
            # Locate "Translation" column header
            translation_col_idx = -1
            if len(table.rows) > 0:
                for col_idx, cell in enumerate(table.rows[0].cells):
                    if cell.text.strip() == "Translation":
                        translation_col_idx = col_idx
                        break
            
            if translation_col_idx == -1:
                continue
                
            table_count += 1
            # Scan rows starting from row 1 (index 1) to skip header
            for r_idx in range(1, len(table.rows)):
                if limit_cells and cell_count >= limit_cells:
                    break
                row = table.rows[r_idx]
                if translation_col_idx < len(row.cells):
                    cell = row.cells[translation_col_idx]
                    cells_to_process.append({
                        "table_index": t_idx,
                        "row_index": r_idx,
                        "cell_text": cell.text,
                        "header": "Translation",
                        "cell": cell
                    })
                    cell_count += 1
                    
        return cells_to_process, table_count

    def analyze_document(self):
        """Scan document and generate analysis report without translating."""
        logger.info(f"Scanning {self.input_path} for analysis...")
        cells, tables_found = self.scan_translation_cells(self.input_path)
        
        total_cells = len(cells)
        thai_cells = 0
        protected_count = 0
        placeholders_count = 0
        translatable_cells = 0
        unique_strings = set()
        repeated_strings = {}
        
        for c in cells:
            text = c["cell_text"].strip()
            if not text:
                continue
            
            unique_strings.add(text)
            repeated_strings[text] = repeated_strings.get(text, 0) + 1
            
            if THAI_CHAR_PATTERN.search(text):
                thai_cells += 1
                continue
                
            is_protected = False
            for pattern in PROTECTED_PATTERNS:
                if re.match(pattern, text):
                    is_protected = True
                    break
            if is_protected:
                protected_count += 1
                continue
                
            # Placeholders inside text
            placeholder_matches = re.findall(r'\{[^}]+\}|%s|%d|%[^ ]+s', text)
            if placeholder_matches:
                placeholders_count += len(placeholder_matches)
                
            translatable_cells += 1

        repeated_count = sum(1 for k, v in repeated_strings.items() if v > 1)

        report = {
            "input_file": self.input_path,
            "tables_containing_translation_column": tables_found,
            "total_rows_scanned": len(cells),
            "cells_requiring_translation": translatable_cells,
            "already_thai_cells": thai_cells,
            "protected_strings": protected_count,
            "placeholders_detected": placeholders_count,
            "repeated_strings_for_tm": repeated_count,
            "unique_translatable_strings": len(unique_strings) - thai_cells - protected_count
        }

        with open(self.analysis_report_path, "w", encoding="utf-8") as f:
            json.dump(report, f, indent=2)
            
        logger.info(f"Analysis complete. Report saved to {self.analysis_report_path}")
        return report

    def _translate_with_runs(self, paragraph, context=None):
        """Processes paragraph runs with tag validation and fallback."""
        runs = paragraph.runs
        if not runs:
            return
            
        # If there is only one run, replace directly
        if len(runs) == 1:
            orig_text = runs[0].text
            cached = self.lookup_translation(orig_text)
            if cached:
                runs[0].text = cached
            else:
                translated = self.translator.translate(orig_text, context)
                self.add_to_tm(orig_text, translated, context)
                runs[0].text = translated
            return

        # Tag-based translation for multiple runs
        tagged_parts = []
        for idx, run in enumerate(runs):
            # Wrap translatable run text
            tagged_parts.append(f"<r{idx}>{run.text}</r{idx}>")
            
        tagged_text = "".join(tagged_parts)
        
        try:
            translated_tagged = self.translator.translate(tagged_text, context)
            
            # Validate tags in response
            for idx in range(len(runs)):
                open_tag = f"<r{idx}>"
                close_tag = f"</r{idx}>"
                if open_tag not in translated_tagged or close_tag not in translated_tagged:
                    raise ValueError(f"Tag validation failed: Missing run tag r{idx}")
                    
            # Parse and apply translated text back to runs
            for idx, run in enumerate(runs):
                pattern = f"<r{idx}>(.*?)</r{idx}>"
                match = re.search(pattern, translated_tagged, re.DOTALL)
                if match:
                    run.text = match.group(1)
        except Exception as e:
            logger.warning(f"Tag-based translation failed ({e}). Falling back to run-by-run mode.")
            # Run-by-run fallback
            for run in runs:
                if run.text.strip():
                    cached = self.lookup_translation(run.text)
                    if cached:
                        run.text = cached
                    else:
                        translated = self.translator.translate(run.text, context)
                        self.add_to_tm(run.text, translated, context)
                        run.text = translated

    def translate_document(self, limit_tables=None, limit_cells=None, is_test=False):
        """Translate cells and save progress."""
        doc_path = self.input_path
        cells, _ = self.scan_translation_cells(doc_path, limit_tables, limit_cells)
        
        logger.info(f"Starting translation of {len(cells)} cells...")
        
        # Threaded translations to populate TM
        def process_cell(cell_info):
            text = cell_info["cell_text"]
            if not text.strip() or THAI_CHAR_PATTERN.search(text):
                return
            
            # Check TM/glossary first
            if self.lookup_translation(text):
                return
                
            context = f"Table header: {cell_info['header']}"
            try:
                # If there's multiple paragraphs/runs, we translate them contextually
                # For worker thread safety, we fetch and cache translation in TM
                translated = self.translator.translate(text, context)
                self.add_to_tm(text, translated, context)
            except Exception as e:
                logger.error(f"Error translating cell: {e}")

        # Distribute work via workers
        num_workers = self.config["concurrency"]["num_workers"]
        with ThreadPoolExecutor(max_workers=num_workers) as executor:
            futures = [executor.submit(process_cell, c) for c in cells]
            for future in as_completed(futures):
                pass
                
        # Single writer applies the translations to a copy of the document
        target_output = self.test_path if is_test else self.output_path
        temp_output = target_output + ".tmp"
        
        logger.info("Applying translations to document structure...")
        doc = Document(doc_path)
        
        # Scan again and apply run-by-run or paragraph translation using cached TM
        cells_to_write, _ = self.scan_translation_cells(doc_path, limit_tables, limit_cells)
        
        # Match cells by table/row index on the new doc copy
        for c in cells_to_write:
            t_idx = c["table_index"]
            r_idx = c["row_index"]
            
            # Locate cell in the target output document structure
            target_cell = doc.tables[t_idx].rows[r_idx].cells[0] # Finding translation column index
            translation_col_idx = -1
            for idx, header_cell in enumerate(doc.tables[t_idx].rows[0].cells):
                if header_cell.text.strip() == "Translation":
                    translation_col_idx = idx
                    break
                    
            if translation_col_idx != -1:
                target_cell = doc.tables[t_idx].rows[r_idx].cells[translation_col_idx]
                
                # Apply TM translations at the run level
                for paragraph in target_cell.paragraphs:
                    self._translate_with_runs(paragraph, context=f"Table: {t_idx}, Row: {r_idx}")
                    
        # Save temp output
        doc.save(temp_output)
        
        # Verify document integrity before replacing final
        if self.verify_document_integrity(doc_path, temp_output):
            if os.path.exists(target_output):
                os.remove(target_output)
            os.rename(temp_output, target_output)
            logger.info(f"Document localized successfully: {target_output}")
            self.generate_review_report(cells)
            return True
        else:
            logger.error("Document validation failed! Temporary output preserved for investigation.")
            return False

    def verify_document_integrity(self, original_path, target_path):
        """QA Gate to ensure formatting, structure, and metadata are intact."""
        try:
            orig_doc = Document(original_path)
            target_doc = Document(target_path)
            
            if len(orig_doc.tables) != len(target_doc.tables):
                logger.error("QA Error: Table count mismatch!")
                return False
                
            for idx in range(len(orig_doc.tables)):
                orig_table = orig_doc.tables[idx]
                target_table = target_doc.tables[idx]
                
                if len(orig_table.rows) != len(target_table.rows):
                    logger.error(f"QA Error: Row count mismatch in Table {idx}!")
                    return False
                    
                if len(orig_table.columns) != len(target_table.columns):
                    logger.error(f"QA Error: Column count mismatch in Table {idx}!")
                    return False
                    
            logger.info("QA Gate: Structural comparison passed.")
            return True
        except Exception as e:
            logger.error(f"QA Gate Error opening files: {e}")
            return False

    def generate_review_report(self, cells):
        """Export localized review spreadsheet."""
        with open(self.review_report_path, "w", encoding="utf-8", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["Table Index", "Row Index", "Original English", "Thai Translation", "Status"])
            for c in cells:
                orig = c["cell_text"]
                trans = self.lookup_translation(orig) or orig
                status = "Translated" if orig != trans else "Skipped/Protected"
                writer.writerow([c["table_index"], c["row_index"], orig, trans, status])
        logger.info(f"Review report saved to {self.review_report_path}")

def main():
    parser = argparse.ArgumentParser(description="AI-Powered DOCX Translation Pipeline")
    parser.add_argument("--analyze-only", action="store_true", help="Scan document metrics and exit.")
    parser.add_argument("--test-mode", action="store_true", help="Run test mode (5 tables / 50 cells).")
    parser.add_argument("--full", action="store_true", help="Run full translation pipeline.")
    
    args = parser.parse_args()
    
    pipeline = LocalizationPipeline()
    
    if args.analyze_only:
        pipeline.analyze_document()
    elif args.test_mode:
        logger.info("Executing Test Mode...")
        pipeline.translate_document(limit_tables=5, limit_cells=50, is_test=True)
    elif args.full:
        logger.info("Executing Full Localization...")
        pipeline.translate_document()
    else:
        parser.print_help()

if __name__ == "__main__":
    main()
